import streamlit as st
import zipfile
import os
import json
import tempfile
import shutil
from pathlib import Path
from docx import Document
from PyPDF2 import PdfReader
import requests
from typing import Dict, List, Tuple

st.set_page_config(page_title="Assignment Grader", page_icon="📝")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from .docx file"""
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from .pdf file"""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from .txt or .md file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with latin-1 encoding as fallback
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def parse_task_document(file_path: str) -> str:
    """Parse task document based on file extension"""
    ext = Path(file_path).suffix.lower()

    if ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.txt', '.md']:
        return extract_text_from_txt(file_path)
    else:
        return ""

def call_openrouter_api(prompt: str, api_key: str, model: str) -> str:
    """Call OpenRouter API"""
    response = requests.post(
        url="https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        },
        json={
            "model": model,
            "messages": [
                {"role": "user", "content": prompt}
            ]
        }
    )

    if response.status_code == 200:
        return response.json()['choices'][0]['message']['content']
    else:
        raise Exception(f"API call failed: {response.status_code} - {response.text}")

def analyze_task_and_create_criteria(task_text: str, api_key: str, model: str) -> Tuple[Dict, str]:
    """Use AI API to analyze task and create grading criteria"""
    prompt = f"""Analyze the following programming task and create a grading criteria.

Task:
{task_text}

Please provide:
1. A JSON object with grading criteria (criteria.json format) with fields like: correctness, code_quality, efficiency, documentation, etc. Each criterion should have a weight and description.
2. A clear description of the task in markdown format (description.md).

Format your response as:
CRITERIA_JSON:
{{json here}}

DESCRIPTION_MD:
{{markdown here}}
"""

    response_text = call_openrouter_api(prompt, api_key, model)

    # Parse the response
    criteria_json = {}
    description_md = ""

    if "CRITERIA_JSON:" in response_text and "DESCRIPTION_MD:" in response_text:
        parts = response_text.split("CRITERIA_JSON:")
        json_part = parts[1].split("DESCRIPTION_MD:")[0].strip()
        description_md = parts[1].split("DESCRIPTION_MD:")[1].strip()

        # Extract JSON from code blocks if present
        if "```json" in json_part:
            json_part = json_part.split("```json")[1].split("```")[0].strip()
        elif "```" in json_part:
            json_part = json_part.split("```")[1].split("```")[0].strip()

        criteria_json = json.loads(json_part)

        # Extract markdown from code blocks if present
        if "```markdown" in description_md:
            description_md = description_md.split("```markdown")[1].split("```")[0].strip()
        elif "```" in description_md:
            description_md = description_md.split("```")[1].split("```")[0].strip()

    return criteria_json, description_md

def grade_submission(submission_code: str, task_description: str, criteria: Dict, api_key: str, model: str) -> Dict:
    """Use AI API to grade a submission"""
    prompt = f"""Grade the following Python submission based on the task description and criteria.

Task Description:
{task_description}

Grading Criteria:
{json.dumps(criteria, indent=2)}

Student Submission:
```python
{submission_code}
```

Please provide a detailed assessment with:
1. Overall score (0-100)
2. Scores for each criterion
3. Detailed feedback
4. Strengths and areas for improvement

Format your response as JSON:
{{
  "overall_score": 85,
  "criteria_scores": {{"criterion_name": score, ...}},
  "feedback": "detailed feedback",
  "strengths": ["strength1", "strength2"],
  "improvements": ["improvement1", "improvement2"]
}}
"""

    response_text = call_openrouter_api(prompt, api_key, model)

    # Extract JSON from response
    if "```json" in response_text:
        json_str = response_text.split("```json")[1].split("```")[0].strip()
    elif "```" in response_text:
        json_str = response_text.split("```")[1].split("```")[0].strip()
    else:
        json_str = response_text.strip()

    return json.loads(json_str)

def create_student_feedback_docx(student_name: str, grade: str, assessment: str,
                                   grading_result: Dict, output_path: str):
    """Create individual student feedback document"""
    doc = Document()

    doc.add_heading(f'Feedback for {student_name}', 0)
    doc.add_paragraph(f'Grade: {grade}')
    doc.add_paragraph(f'Assessment: {assessment}')
    doc.add_paragraph(f'Overall Score: {grading_result["overall_score"]}/100')

    doc.add_heading('Detailed Feedback', level=1)
    doc.add_paragraph(grading_result['feedback'])

    doc.add_heading('Strengths', level=1)
    for strength in grading_result['strengths']:
        doc.add_paragraph(strength, style='List Bullet')

    doc.add_heading('Areas for Improvement', level=1)
    for improvement in grading_result['improvements']:
        doc.add_paragraph(improvement, style='List Bullet')

    doc.add_heading('Criteria Scores', level=1)
    for criterion, score in grading_result['criteria_scores'].items():
        doc.add_paragraph(f'{criterion}: {score}')

    doc.save(output_path)

def create_teacher_report_docx(all_results: List[Dict], output_path: str):
    """Create comprehensive teacher report"""
    doc = Document()

    doc.add_heading('Assessment Grade Report', 0)
    doc.add_paragraph(f'Total Students: {len(all_results)}')

    avg_score = sum(r['grading_result']['overall_score'] for r in all_results) / len(all_results)
    doc.add_paragraph(f'Average Score: {avg_score:.2f}/100')

    doc.add_heading('Individual Results', level=1)

    for result in sorted(all_results, key=lambda x: x['grading_result']['overall_score'], reverse=True):
        doc.add_heading(f"{result['name']} - {result['grade']} - {result['assessment']}", level=2)
        doc.add_paragraph(f"Score: {result['grading_result']['overall_score']}/100")
        doc.add_paragraph(f"Feedback: {result['grading_result']['feedback'][:200]}...")
        doc.add_paragraph('')

    doc.save(output_path)

def process_zip_file(zip_file, api_key: str, model: str, progress_bar, status_text):
    """Main processing function"""
    with tempfile.TemporaryDirectory() as tmpdir:
        # Extract zip file
        zip_path = os.path.join(tmpdir, "upload.zip")
        with open(zip_path, "wb") as f:
            f.write(zip_file.getvalue())

        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(tmpdir)

        # Find task and submissions folders
        task_folder = None
        submissions_folder = None

        for root, dirs, files in os.walk(tmpdir):
            for d in dirs:
                if d.lower() == 'task' and not task_folder:
                    task_folder = os.path.join(root, d)
                if d.lower() == 'submissions' and not submissions_folder:
                    submissions_folder = os.path.join(root, d)

        if not task_folder or not submissions_folder:
            # Show debug info
            all_folders = []
            for root, dirs, files in os.walk(tmpdir):
                for d in dirs:
                    all_folders.append(os.path.join(root, d))
            st.error(f"Could not find 'task' and 'submissions' folders in the zip file.\n\nFolders found: {all_folders}")
            return None

        status_text.text("Processing task documents...")

        # Check for existing criteria.json and description.md
        criteria_path = os.path.join(task_folder, 'criteria.json')
        description_path = os.path.join(task_folder, 'description.md')

        if os.path.exists(criteria_path) and os.path.exists(description_path):
            with open(criteria_path, 'r') as f:
                criteria = json.load(f)
            with open(description_path, 'r') as f:
                task_description = f.read()
        else:
            # Parse all task documents
            task_text = ""
            for file in os.listdir(task_folder):
                file_path = os.path.join(task_folder, file)
                if os.path.isfile(file_path):
                    task_text += parse_task_document(file_path) + "\n\n"

            # Generate criteria and description
            criteria, task_description = analyze_task_and_create_criteria(task_text, api_key, model)

        progress_bar.progress(0.2)

        # Process submissions - handle .py files, .zip files, and folders
        submissions = []
        submission_paths = {}  # Maps filename to full path
        all_files = os.listdir(submissions_folder)

        status_text.text("Extracting submissions...")

        for item in all_files:
            item_path = os.path.join(submissions_folder, item)

            # Skip macOS metadata files
            if item.startswith('.') or item.startswith('._'):
                continue

            # Handle .py files directly in submissions folder
            if item.endswith('.py') and os.path.isfile(item_path):
                submissions.append(item)
                submission_paths[item] = item_path

            # Handle .zip files in submissions folder
            elif item.endswith('.zip') and os.path.isfile(item_path):
                # Extract the zip file
                extract_dir = os.path.join(submissions_folder, f"_extracted_{item[:-4]}")
                os.makedirs(extract_dir, exist_ok=True)

                try:
                    with zipfile.ZipFile(item_path, 'r') as zip_ref:
                        zip_ref.extractall(extract_dir)

                    # Find .py files in the extracted content
                    for root, dirs, files in os.walk(extract_dir):
                        for file in files:
                            if file.endswith('.py') and not file.startswith('.') and not file.startswith('._'):
                                # Use original naming from zip filename
                                py_file_path = os.path.join(root, file)
                                # Try to extract student info from zip name
                                zip_base = item[:-4]  # Remove .zip
                                submissions.append(f"{zip_base}.py")
                                submission_paths[f"{zip_base}.py"] = py_file_path
                                break  # Take first .py file found
                        if f"{zip_base}.py" in submissions:
                            break
                except Exception as e:
                    status_text.text(f"Warning: Could not extract {item}: {str(e)}")
                    continue

            # Handle folders in submissions folder
            elif os.path.isdir(item_path):
                # Look for .py files in the folder
                for root, dirs, files in os.walk(item_path):
                    for file in files:
                        if file.endswith('.py') and not file.startswith('.') and not file.startswith('._'):
                            py_file_path = os.path.join(root, file)
                            # Use folder name for student identification
                            submissions.append(f"{item}.py")
                            submission_paths[f"{item}.py"] = py_file_path
                            break  # Take first .py file found
                    if f"{item}.py" in submissions:
                        break

        if not submissions:
            st.error(f"No Python submissions found in the submissions folder.\n\nFiles/folders found: {all_files}")
            return None

        all_results = []
        output_folder = os.path.join(tmpdir, 'output')
        os.makedirs(output_folder, exist_ok=True)

        for idx, submission_file in enumerate(submissions):
            status_text.text(f"Grading {submission_file}...")

            # Parse filename: name_surname_grade_assessment.py
            parts = submission_file.replace('.py', '').split('_')
            if len(parts) >= 4:
                name = parts[0]
                surname = parts[1]
                grade = parts[2]
                assessment = '_'.join(parts[3:])
            else:
                name = "Unknown"
                surname = ""
                grade = ""
                assessment = ""

            # Read submission code with fallback encoding
            submission_path = submission_paths[submission_file]
            try:
                with open(submission_path, 'r', encoding='utf-8') as f:
                    submission_code = f.read()
            except UnicodeDecodeError:
                # Try with latin-1 encoding as fallback
                try:
                    with open(submission_path, 'r', encoding='latin-1') as f:
                        submission_code = f.read()
                except:
                    status_text.text(f"Skipping {submission_file} - unable to read file")
                    continue
            except Exception as e:
                status_text.text(f"Skipping {submission_file} - error: {str(e)}")
                continue

            # Check if file contains readable Python code (not binary)
            if submission_code.strip() == '' or len([c for c in submission_code[:100] if ord(c) < 32 and c not in '\n\r\t']) > 10:
                status_text.text(f"Skipping {submission_file} - appears to be binary or empty")
                continue

            # Grade submission
            grading_result = grade_submission(submission_code, task_description, criteria, api_key, model)

            # Create student feedback document
            student_name = f"{name}_{surname}"
            feedback_filename = f"{name}_{surname}_{grade}_{assessment}_feedback.docx"
            feedback_path = os.path.join(output_folder, feedback_filename)
            create_student_feedback_docx(student_name, grade, assessment, grading_result, feedback_path)

            all_results.append({
                'name': student_name,
                'grade': grade,
                'assessment': assessment,
                'grading_result': grading_result
            })

            progress_bar.progress(0.2 + (0.7 * (idx + 1) / len(submissions)))

        status_text.text("Creating teacher report...")

        # Create teacher report
        report_path = os.path.join(output_folder, 'assessment_grade_report.docx')
        create_teacher_report_docx(all_results, report_path)

        progress_bar.progress(1.0)
        status_text.text("Done!")

        # Read all files into memory before temp directory is deleted
        files_data = {}
        for file in os.listdir(output_folder):
            file_path = os.path.join(output_folder, file)
            with open(file_path, 'rb') as f:
                files_data[file] = f.read()

        return files_data

# Available models
MODELS = {
    "Free Models": {
        "Xiaomi Mimo V2 Flash (Free)": "xiaomi/mimo-v2-flash:free",
        "Mistral Devstral 2512 (Free)": "mistralai/devstral-2512:free",
    },
    "Premium Models": {
        "Claude Sonnet 3.5": "anthropic/claude-3.5-sonnet",
        "GPT-4 Turbo": "openai/gpt-4-turbo",
        "GPT-4o": "openai/gpt-4o",
        "Claude Opus 3": "anthropic/claude-3-opus",
    }
}

# Initialize session state for storing files
if 'graded_files' not in st.session_state:
    st.session_state.graded_files = None

# Streamlit UI
st.title("📝 Assignment Grader")
st.write("Upload a .zip file containing 'task' and 'submissions' folders to automatically grade student assignments.")

# API Key input
api_key = st.text_input("Enter your OpenRouter API Key:", type="password")
st.markdown("Get your free API key at [OpenRouter](https://openrouter.ai/)")

# Model selection
st.subheader("Select AI Model")
model_category = st.radio("Model Type:", ["Free Models", "Premium Models"])
model_name = st.selectbox("Choose Model:", list(MODELS[model_category].keys()))
selected_model = MODELS[model_category][model_name]

if model_category == "Free Models":
    st.info("You selected a FREE model - no costs will be incurred!")
else:
    st.warning("Premium models incur costs per token used.")

# File upload
uploaded_file = st.file_uploader("Upload ZIP file", type=['zip'])

# Grade button
if st.button("Grade Submissions", type="primary"):
    if not api_key:
        st.error("Please enter your OpenRouter API key")
    elif not uploaded_file:
        st.error("Please upload a ZIP file")
    else:
        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            files_data = process_zip_file(uploaded_file, api_key, selected_model, progress_bar, status_text)

            if files_data:
                st.session_state.graded_files = files_data
                st.success("Grading completed successfully!")
                st.rerun()
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            import traceback
            st.code(traceback.format_exc())

# Display download buttons if files are available
if st.session_state.graded_files:
    st.subheader("Download Results")

    # Create a ZIP file containing all results
    import io
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename, file_data in st.session_state.graded_files.items():
            zip_file.writestr(filename, file_data)

    zip_buffer.seek(0)

    # Download All button
    st.download_button(
        label="📦 Download All Files (ZIP)",
        data=zip_buffer,
        file_name="grading_results.zip",
        mime="application/zip",
        type="primary"
    )

    st.markdown("---")
    st.markdown("**Individual Files:**")

    # Individual download buttons
    for filename, file_data in st.session_state.graded_files.items():
        st.download_button(
            label=f"📄 {filename}",
            data=file_data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{filename}"
        )

    # Clear results button
    if st.button("🗑️ Clear Results", type="secondary"):
        st.session_state.graded_files = None
        st.rerun()

st.markdown("---")
st.markdown("### Instructions")
st.markdown("""
1. Get a free API key from [OpenRouter](https://openrouter.ai/)
2. Prepare a ZIP file with two folders:
   - `task`: Contains task description files (.docx, .pdf, .md, or .txt)
   - `submissions`: Contains student Python files named as `name_surname_grade_assessment.py`
3. Enter your OpenRouter API key
4. Select a model (use Free Models for zero cost!)
5. Upload the ZIP file
6. Click 'Grade Submissions'
7. Download the generated feedback files and teacher report
""")
